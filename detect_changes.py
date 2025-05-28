import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from googletrans import Translator
import difflib

# ========== Step 1: Extract tracked changes from English docx ==========
def extract_tracked_changes(docx_path):
    changes = []
    with zipfile.ZipFile(docx_path) as docx_zip:
        with docx_zip.open("word/document.xml") as document_xml:
            tree = ET.parse(document_xml)
            root = tree.getroot()
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            for para_index, para in enumerate(root.findall('.//w:p', namespaces)):
                inserted = para.findall('.//w:ins//w:t', namespaces)
                deleted = para.findall('.//w:del//w:t', namespaces)
                if inserted:
                    text = ''.join([t.text for t in inserted if t.text])
                    if text.strip():
                        changes.append({
                            'paragraph': para_index,
                            'type': 'insert',
                            'text': text.strip()
                        })
                if deleted:
                    text = ''.join([t.text for t in deleted if t.text])
                    if text.strip():
                        changes.append({
                            'paragraph': para_index,
                            'type': 'delete',
                            'text': text.strip()
                        })
    return changes

# ========== Step 2: Translate English to Chinese ==========
def translate_to_chinese(texts):
    translator = Translator()
    translations = []
    for text in texts:
        try:
            translated = translator.translate(text, src='en', dest='zh-cn').text
            translations.append(translated)
        except Exception as e:
            translations.append(f"Error translating: {text}")
    return translations

# ========== Step 3: Load Chinese paragraphs ==========
def load_chinese_paragraphs(docx_path):
    doc = Document(docx_path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

# ========== Step 4: Match translated Chinese to Chinese doc ==========
def match_translations_to_chinese(translations, chinese_paragraphs):
    matches = []
    for translation in translations:
        best_match = None
        best_ratio = 0
        best_index = -1
        for i, para in enumerate(chinese_paragraphs):
            ratio = difflib.SequenceMatcher(None, translation, para).ratio()
            if ratio > best_ratio:
                best_ratio = ratio
                best_match = para
                best_index = i
        matches.append({
            'translated_text': translation,
            'matched_paragraph': best_match,
            'paragraph_index': best_index,
            'similarity': round(best_ratio, 2)
        })
    return matches

# ========== Step 5: Insert tracker text into Chinese document ==========
def insert_tracker_text(doc, report):
    for item in report:
        para_index = item['paragraph_index']
        if para_index < 0 or para_index >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[para_index]
        # Append tracking note as italic text
        note = f"\n[TRACK-{item['change_type'].upper()}: {item['translated_chinese_text']}]"
        para.add_run(note).italic = True

# ========== Main Pipeline ==========
def main():
    # Input paths
    english_docx = "edited_en.docx"
    chinese_docx = "original_cn.docx"
    output_docx = "Chinese_with_Tracked_Changes.docx"

    # Step 1: Extract changes
    print("🔍 Extracting tracked changes from English document...")
    changes = extract_tracked_changes(english_docx)
    english_changes = [c['text'] for c in changes]

    # Step 2: Translate changes
    print("🌐 Translating changes to Chinese...")
    translations = translate_to_chinese(english_changes)

    # Step 3: Load Chinese doc paragraphs
    print("📄 Loading Chinese document...")
    chinese_paragraphs = load_chinese_paragraphs(chinese_docx)

    # Step 4: Match translated changes to Chinese
    print("🧠 Matching translations to Chinese paragraphs...")
    matches = match_translations_to_chinese(translations, chinese_paragraphs)

    # Step 5: Combine into report
    final_report = []
    for change, match in zip(changes, matches):
        final_report.append({
            'change_type': change['type'],
            'original_english_text': change['text'],
            'translated_chinese_text': match['translated_text'],
            'matched_chinese_paragraph': match['matched_paragraph'],
            'paragraph_index': match['paragraph_index'],
            'similarity': match['similarity']
        })

    # Step 6: Inject tracker into Chinese doc
    print("✏️ Writing changes into Chinese document...")
    doc = Document(chinese_docx)
    insert_tracker_text(doc, final_report)
    doc.save(output_docx)

    print(f"✅ Done! Output saved as: {output_docx}")

if __name__ == "__main__":
    main()
